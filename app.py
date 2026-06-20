from __future__ import annotations

import io
import os
import re
import subprocess
import sys
import threading
import textwrap
import time
from pathlib import Path

import requests
import streamlit as st


MODEL_URL = "http://0.0.0.0:8000/chat/local_llm/"
APP_HOST = "0.0.0.0"
APP_PORT = 8550
APP_URL = f"http://127.0.0.1:{APP_PORT}"
OUTPUT_DIR = Path(__file__).resolve().parent / "translations"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

SUPPORTED_TYPES = ["txt", "md", "docx", "pdf"]
CHUNK_SIZE = 3500


class TranslationError(RuntimeError):
    pass


def launch_streamlit() -> None:
    env = os.environ.copy()
    env["STREAMLIT_LAUNCHED_BY_APP"] = "1"

    threading.Thread(target=open_edge_after_delay, daemon=True).start()
    subprocess.run(
        [
            sys.executable,
            "-m",
            "streamlit",
            "run",
            str(Path(__file__).resolve()),
            "--server.address",
            APP_HOST,
            "--server.port",
            str(APP_PORT),
            "--server.headless",
            "true",
        ],
        env=env,
        check=False,
    )


def open_edge_after_delay() -> None:
    time.sleep(2)
    open_windows_edge(APP_URL)


def open_windows_edge(url: str) -> None:
    try:
        subprocess.Popen(
            ["cmd.exe", "/c", "start", "", "msedge", url],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
    except OSError:
        print(f"Open this URL in Edge: {url}")


def normalize_endpoint(value: str) -> str:
    endpoint = value.strip() or MODEL_URL
    if not endpoint.endswith("/"):
        endpoint += "/"
    return endpoint


def clean_filename(name: str) -> str:
    stem = Path(name).stem or "document"
    safe = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._-")
    return safe or "document"


def extract_text_from_file(file_name: str, data: bytes) -> str:
    suffix = Path(file_name).suffix.lower()
    if suffix in {".txt", ".md"}:
        return data.decode("utf-8", errors="replace")
    if suffix == ".docx":
        return extract_docx_text(data)
    if suffix == ".pdf":
        return extract_pdf_text(data)
    raise TranslationError(f"Unsupported file type: {suffix or 'unknown'}")


def extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise TranslationError("Install python-docx to translate .docx files.") from exc

    document = Document(io.BytesIO(data))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise TranslationError("Install pypdf to translate .pdf files.") from exc

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(page.strip() for page in pages if page.strip())


def split_text(text: str, max_chars: int = CHUNK_SIZE) -> list[str]:
    blocks = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
    chunks: list[str] = []
    current: list[str] = []
    current_size = 0

    for block in blocks:
        if len(block) > max_chars:
            if current:
                chunks.append("\n\n".join(current))
                current = []
                current_size = 0
            chunks.extend(textwrap.wrap(block, width=max_chars, break_long_words=False, replace_whitespace=False))
            continue

        next_size = current_size + len(block) + (2 if current else 0)
        if current and next_size > max_chars:
            chunks.append("\n\n".join(current))
            current = [block]
            current_size = len(block)
        else:
            current.append(block)
            current_size = next_size

    if current:
        chunks.append("\n\n".join(current))

    return chunks or [text.strip()]


def translate_chunk(endpoint: str, chunk: str, chunk_number: int, total_chunks: int) -> str:
    system_prompt = (
        "You are a professional translator. Translate Spanish text into natural English. "
        "Preserve paragraph breaks, names, numbers, lists, and document meaning. "
        "Return only the English translation.")
    # system_prompt = "Translate this Spanish document chunk to English."   # DON'T DELETE

    user_prompt = (
        f"Chunk {chunk_number} of {total_chunks}:\n\n{chunk}"
    )
    request_body = {
        "req_id": f"translate_{chunk_number:04d}",
        "query": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }

    response = requests.post(endpoint, json=request_body, timeout=180)
    response.raise_for_status()
    payload = response.json()

    if not payload.get("success", False):
        message = payload.get("message") or payload.get("error") or "Model request failed."
        raise TranslationError(str(message))

    translated = payload.get("response")
    if not isinstance(translated, str) or not translated.strip():
        raise TranslationError("Model returned an empty translation.")
    return translated.strip()


def translate_document(endpoint: str, text: str) -> str:
    chunks = split_text(text)
    progress = st.progress(0)
    status = st.empty()
    translations: list[str] = []

    for index, chunk in enumerate(chunks, start=1):
        status.info(f"Translating chunk {index} of {len(chunks)}")
        progress.progress((index - 1) / len(chunks))
        translations.append(translate_chunk(endpoint, chunk, index, len(chunks)))

    progress.progress(1.0)
    status.success("Translation complete")
    return "\n\n".join(translations)


def save_translation(source_name: str, translated_text: str) -> Path:
    output_path = OUTPUT_DIR / f"{clean_filename(source_name)}_english.txt"
    output_path.write_text(translated_text, encoding="utf-8")
    return output_path


def render_styles() -> None:
    st.markdown(
        """
        <style>
        :root {
            --page: #f6f4e8;
            --ink: #283044;
            --muted: #637084;
            --accent: #ca6180;
            --accent-dark: #a84663;
        }
        .stApp {
            background: var(--page);
            color: var(--ink);
        }
        .main .block-container {
            max-width: 1240px;
            padding-top: 2rem;
            padding-bottom: 3rem;
        }
        section[data-testid="stSidebar"] {
            background: linear-gradient(180deg, #7daacb, #d7e8f2);
            border-right: 1px solid #5f91b6;
        }
        .hero-band {
            border: 1px solid #744577;
            border-radius: 8px;
            background: #EFD2B0;
            padding: 1.2rem 1.35rem;
            margin-bottom: 1rem;
        }
        .hero-title {
            margin: 0 0 .25rem 0;
            color: #E13F7C !important;
            font-size: 2rem;
            font-weight: 800;
            line-height: 1.15;
            letter-spacing: 0;
        }
        .hero-band p {
            margin: 0;
            color: var(--muted);
            font-size: 1rem;
        }
        .section-label {
            color: var(--ink);
            font-size: 1.05rem;
            font-weight: 700;
            margin: .5rem 0 .55rem 0;
        }
        .status-strip {
            border: 1px solid #b89480;
            border-radius: 8px;
            background: #F3E3D0;
            padding: .75rem .9rem;
            margin: .6rem 0 1rem 0;
        }
        .metric-name {
            color: var(--muted);
            font-size: .78rem;
            font-weight: 700;
            text-transform: uppercase;
        }
        .metric-value {
            color: var(--ink);
            font-size: 1rem;
            font-weight: 650;
            margin-top: .2rem;
            overflow-wrap: anywhere;
        }
        div[data-testid="stFileUploader"] {
            border: 1px dashed #cade80;
            border-radius: 8px;
            background: rgba(158, 211, 220, .20);
            padding: .5rem .75rem;
        }
        div[data-testid="stTextArea"] textarea {
            border-radius: 8px;
            border-color: #9ed3dc;
            background-color: #f7fdff;
            color: #2b2f3a;
            font-family: "Segoe UI", Calibri, sans-serif;
            line-height: 1.45;
        }
        .stButton > button,
        .stDownloadButton > button {
            border-radius: 7px;
            display: block;
            min-height: 2.65rem;
            font-weight: 700;
            color: #0f172a !important;
            border: 1px solid #a84663;
            width: 100%;
        }
        div.stButton,
        div.stDownloadButton {
            max-width: 260px;
            margin-left: auto;
            margin-right: auto;
        }
        .stButton > button[kind="primary"] {
            background: var(--accent);
            border-color: var(--accent-dark);
            color: #ffffff !important;
        }
        .stButton > button[kind="primary"]:hover {
            background: var(--accent-dark);
            border-color: var(--accent-dark);
        }
        .stDownloadButton > button {
            background: #fefd99;
            border-color: #d6cf55;
            color: #283044 !important;
        }
        .stDownloadButton > button:hover {
            background: #f6ed72;
            border-color: #bdb64b;
            color: #1f2937 !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def init_session_state() -> None:
    defaults = {
        "current_file_name": "",
        "current_source_text": "",
        "current_translation": "",
    }
    for key, value in defaults.items():
        st.session_state.setdefault(key, value)


def clear_document_state() -> None:
    st.session_state["current_file_name"] = ""
    st.session_state["current_source_text"] = ""
    st.session_state["current_translation"] = ""


def render_app() -> None:
    st.set_page_config(
        page_title="Spanish to English Translator",
        page_icon="TR",
        layout="wide",
        initial_sidebar_state="expanded",
    )
    render_styles()
    init_session_state()

    st.markdown(
        """
        <div class="hero-band">
            <div class="hero-title">Spanish to English Translator</div>
            <p>Upload a document, inspect the extracted Spanish text, and translate it through the local Gemma4 endpoint.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.sidebar:
        st.subheader("Translation Settings")
        endpoint = st.text_input("Model endpoint", value=MODEL_URL)
        st.caption("Output folder")
        st.code(str(OUTPUT_DIR), language=None)

    st.markdown('<div class="section-label">Document</div>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader(
        "Choose or drag a Spanish document",
        type=SUPPORTED_TYPES,
        accept_multiple_files=False,
        label_visibility="collapsed",
    )

    if uploaded_file:
        if uploaded_file.name != st.session_state["current_file_name"]:
            try:
                source_text = extract_text_from_file(uploaded_file.name, uploaded_file.getvalue())
            except TranslationError as exc:
                st.error(str(exc))
                return

            if not source_text.strip():
                st.error("No readable text was found in this document.")
                return

            st.session_state["current_file_name"] = uploaded_file.name
            st.session_state["current_source_text"] = source_text
            st.session_state["current_translation"] = ""
    elif st.session_state["current_source_text"]:
        clear_document_state()

    if not st.session_state["current_source_text"]:
        st.info("Supported formats: TXT, MD, DOCX, PDF")
        return

    source_text = st.session_state["current_source_text"]

    st.markdown(
        f"""
        <div class="status-strip">
            <div class="metric-name">Loaded document</div>
            <div class="metric-value">{st.session_state["current_file_name"]}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    left, right = st.columns(2)
    with left:
        st.markdown('<div class="section-label">Spanish text preview</div>', unsafe_allow_html=True)
        st.text_area(
            "Source",
            value=source_text[:6000],
            height=340,
            label_visibility="collapsed",
        )

    with right:
        st.markdown('<div class="section-label">English translation</div>', unsafe_allow_html=True)
        if st.session_state["current_translation"]:
            st.text_area(
                "Translation",
                value=st.session_state["current_translation"],
                height=340,
                label_visibility="collapsed",
            )
        else:
            st.info("Translated text will appear here.")

    action_left, action_right = st.columns(2)

    with action_left:
        translate_clicked = st.button("Translate", type="primary")

    with action_right:
        if st.session_state["current_translation"]:
            st.download_button(
                "Download",
                data=st.session_state["current_translation"].encode("utf-8"),
                file_name=f"{clean_filename(st.session_state['current_file_name'])}_english.txt",
                mime="text/plain",
            )

    if translate_clicked:
        try:
            translated_text = translate_document(
                normalize_endpoint(endpoint),
                st.session_state["current_source_text"],
            )
            output_path = save_translation(st.session_state["current_file_name"], translated_text)
        except (requests.RequestException, TranslationError) as exc:
            st.error(f"Translation failed: {exc}")
            return

        st.session_state["current_translation"] = translated_text
        st.success(f"Saved translation to {output_path}")
        st.rerun()


if __name__ == "__main__" and os.environ.get("STREAMLIT_LAUNCHED_BY_APP") != "1":
    launch_streamlit()
else:
    render_app()
